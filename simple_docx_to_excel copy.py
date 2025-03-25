import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import subprocess
import platform
from docx import Document
import openpyxl
from openpyxl.utils import get_column_letter
import re
from datetime import datetime

class SimpleDocxToExcelApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Таблицы из DOCX в Excel")
        self.root.geometry("600x400")
        
        # Переменные для хранения путей к файлам
        self.docx_path = None
        self.excel_path = None
        
        # Создание интерфейса
        self.create_gui()
    
    def create_gui(self):
        """Создание простого интерфейса"""
        # Главный фрейм
        main_frame = ttk.Frame(self.root, padding=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Заголовок
        title_label = ttk.Label(main_frame, text="Конвертация таблиц из DOCX в Excel", font=("Arial", 14, "bold"))
        title_label.pack(pady=10)
        
        # Фрейм для выбора файла
        file_frame = ttk.LabelFrame(main_frame, text="Выбор DOCX файла", padding=10)
        file_frame.pack(fill=tk.X, pady=10)
        
        # Поле для отображения пути к файлу
        self.file_path_var = tk.StringVar()
        file_path_entry = ttk.Entry(file_frame, textvariable=self.file_path_var, width=50)
        file_path_entry.grid(row=0, column=0, padx=5, pady=5, sticky="ew")
        
        # Кнопка выбора файла
        browse_button = ttk.Button(file_frame, text="Обзор...", command=self.select_file)
        browse_button.grid(row=0, column=1, padx=5, pady=5)
        
        file_frame.columnconfigure(0, weight=1)
        
        # Информационное поле
        info_frame = ttk.LabelFrame(main_frame, text="Статус", padding=10)
        info_frame.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Текстовое поле для статуса
        self.status_text = tk.Text(info_frame, wrap=tk.WORD, height=8)
        self.status_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Полоса прокрутки
        scrollbar = ttk.Scrollbar(info_frame, orient="vertical", command=self.status_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.status_text.configure(yscrollcommand=scrollbar.set)
        
        self.update_status("Выберите DOCX файл для начала обработки.")
        
        # Фрейм для кнопок
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill=tk.X, pady=10)
        
        # Кнопка обработки
        process_button = ttk.Button(
            button_frame, 
            text="Обработать (конвертировать + удалить столбцы A и C)", 
            command=self.process_file
        )
        process_button.pack(side=tk.LEFT, padx=5)
        
        # Кнопка выхода
        exit_button = ttk.Button(button_frame, text="Выход", command=self.root.destroy)
        exit_button.pack(side=tk.RIGHT, padx=5)
    
    def update_status(self, message):
        """Обновление статуса в текстовом поле"""
        self.status_text.config(state=tk.NORMAL)
        self.status_text.delete(1.0, tk.END)
        self.status_text.insert(tk.END, message)
        self.status_text.config(state=tk.DISABLED)
        self.root.update_idletasks()
    
    def select_file(self):
        """Выбор DOCX файла"""
        file_path = filedialog.askopenfilename(
            title="Выберите DOCX файл",
            filetypes=[("DOCX файлы", "*.docx"), ("Все файлы", "*.*")]
        )
        if file_path:
            self.docx_path = file_path
            self.file_path_var.set(file_path)
            
            # Создаем путь для сохранения Excel
            base_name = os.path.basename(file_path)
            name, _ = os.path.splitext(base_name)
            self.excel_path = os.path.join(os.path.dirname(file_path), f"{name}.xlsx")
            
            self.update_status(f"Выбран файл: {file_path}\n\nНажмите кнопку 'Обработать' для конвертации таблиц и удаления столбцов A и C.")
    
    def process_file(self):
        """Обработка файла: конвертация DOCX в Excel и удаление столбцов A и C"""
        if not self.docx_path:
            messagebox.showerror("Ошибка", "Сначала выберите DOCX файл")
            return
        
        try:
            self.update_status("Обработка файла...\nШаг 1: Извлечение таблиц из DOCX...\n\nПравила обработки:\n- Все таблицы из Word перенесутся в Excel\n- Первая строка удаляется, если во второй ячейке НЕТ даты\n- Первая строка сохраняется, если во второй ячейке ЕСТЬ дата\n- Столбцы A и C будут удалены\n- Даты во втором столбце будут приведены к формату ДД.ММ.ГГГГ")
            
            # Шаг 1: Извлечение таблиц из DOCX в Excel
            table_count = self.convert_docx_to_excel(self.docx_path, self.excel_path)
            
            if table_count > 0:
                self.update_status(f"Шаг 1: Таблицы успешно извлечены ({table_count} шт.)\n\nШаг 2: Удаление первых строк и столбцов A и C...")
                
                # Шаг 2: Удаление столбцов A и C, проверка и удаление первой строки
                sheets_processed, rows_deleted, dates_normalized = self.remove_columns(self.excel_path)
                
                # Финальное сообщение
                self.update_status(
                    f"Обработка успешно завершена!\n\n"
                    f"- Извлечено таблиц: {table_count}\n"
                    f"- Обработано листов: {sheets_processed}\n"
                    f"- Удалено первых строк: {rows_deleted}\n"
                    f"- Нормализовано дат: {dates_normalized}\n"
                    f"- Удалены столбцы: A и C\n\n"
                    f"Результат сохранен в: {self.excel_path}"
                )
                
                # Открываем Excel-файл
                self.open_file(self.excel_path)
                
                messagebox.showinfo("Успех", f"Обработка завершена. Таблицы сохранены в {self.excel_path}")
            else:
                self.update_status("В документе не найдено таблиц.")
                messagebox.showwarning("Предупреждение", "В документе не найдено таблиц")
        
        except Exception as e:
            self.update_status(f"Произошла ошибка при обработке файла:\n{str(e)}")
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")
    
    def convert_docx_to_excel(self, docx_path, excel_path):
        """Извлечение таблиц из DOCX и сохранение в Excel"""
        # Открываем DOCX-файл
        document = Document(docx_path)
        
        # Проверяем, есть ли таблицы
        if not document.tables:
            return 0
        
        # Создаем новую рабочую книгу Excel
        workbook = openpyxl.Workbook()
        # Удаляем стандартный лист
        default_sheet = workbook.active
        workbook.remove(default_sheet)
        
        # Счетчик таблиц
        table_count = 0
        
        # Для каждой таблицы из docx
        for i, table in enumerate(document.tables):
            # Создаем новый лист для каждой таблицы
            sheet_name = f"Таблица_{i+1}"
            sheet = workbook.create_sheet(title=sheet_name)
            table_count += 1
            
            # Копируем данные из таблицы docx в Excel
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    # Excel использует индексацию с 1, а не с 0
                    excel_row = row_idx + 1
                    excel_col = col_idx + 1
                    sheet.cell(row=excel_row, column=excel_col).value = cell.text
            
            # Автоподбор ширины столбцов
            self._adjust_column_width(sheet)
        
        # Сохраняем Excel-файл
        workbook.save(excel_path)
        
        return table_count
    
    def remove_columns(self, excel_path):
        """Удаление столбцов A и C из Excel-файла и обработка первой строки"""
        # Загружаем рабочую книгу
        workbook = openpyxl.load_workbook(excel_path)
        
        # Колонки для удаления в обратном порядке (C, A)
        # Важно: удаляем сначала большие индексы, потом меньшие,
        # чтобы не смещались индексы колонок при удалении
        columns_to_remove = [3, 1]  # C = 3, A = 1
        
        sheets_processed = 0
        rows_deleted = 0
        dates_normalized = 0
        
        # Обрабатываем каждый лист
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheets_processed += 1
            
            # ВАЖНО: Сначала проверяем, нужно ли удалить первую строку
            # Получаем значение ВТОРОЙ ячейки (B1) для проверки
            second_cell_value = sheet.cell(row=1, column=2).value
            
            # Определяем, нужно ли удалять первую строку
            delete_first_row = not self._is_date(second_cell_value)
            
            # Удаляем столбцы
            for col_idx in columns_to_remove:
                sheet.delete_cols(col_idx, 1)
            
            # Теперь удаляем первую строку, если нужно
            if delete_first_row:
                sheet.delete_rows(1, 1)
                rows_deleted += 1
            
            # Нормализуем даты в первом столбце (бывший B, теперь A после удаления)
            normalized_count = self._normalize_dates(sheet)
            dates_normalized += normalized_count
            
            # Автоподбор ширины столбцов
            self._adjust_column_width(sheet)
        
        # Сохраняем изменения
        workbook.save(excel_path)
        
        return sheets_processed, rows_deleted, dates_normalized
    
    def _normalize_dates(self, sheet):
        """
        Нормализует даты в первом столбце (бывший столбец B) к формату ДД.ММ.ГГГГ
        
        Возвращает количество нормализованных дат
        """
        normalized_count = 0
        column_index = 1  # Первый столбец (A после удаления)
        
        # Обрабатываем все ячейки в первом столбце
        for row in range(1, sheet.max_row + 1):
            cell = sheet.cell(row=row, column=column_index)
            value = cell.value
            
            # Пропускаем пустые ячейки
            if not value:
                continue
                
            value_str = str(value).strip()
            
            # Проверяем различные форматы дат и преобразуем их
            normalized_date = self._parse_and_normalize_date(value_str)
            
            if normalized_date:
                cell.value = normalized_date
                normalized_count += 1
        
        return normalized_count
    
    def _parse_and_normalize_date(self, date_str):
            """
            Парсит различные форматы дат и нормализует их к формату ДД.ММ.ГГГГ
            
            Поддерживает:
            - ДД.ММ.ГГ -> ДД.ММ.ГГГГ
            - ДД.ММ.ГГГГ (оставляем как есть)
            - ДДММ.ГГ -> ДД.ММ.ГГГГ (пропущена точка)
            - ДД/ММ/ГГ -> ДД.ММ.ГГГГ
            - ДД-ММ-ГГ -> ДД.ММ.ГГГГ
            - ДДММГГГГ -> ДД.ММ.ГГГГ (без разделителей)
            """
            # Проверяем различные форматы даты
            
            # Формат ДД.ММ.ГГ (двузначный год)
            match = re.match(r'^(\d{1,2})\.(\d{1,2})\.(\d{2})$', date_str)
            if match:
                day, month, year = match.groups()
                full_year = self._expand_year(year)
                return f"{int(day):02d}.{int(month):02d}.{full_year}"
            
            # Формат ДД.ММ.ГГГГ (четырехзначный год)
            match = re.match(r'^(\d{1,2})\.(\d{1,2})\.(\d{4})$', date_str)
            if match:
                day, month, year = match.groups()
                return f"{int(day):02d}.{int(month):02d}.{year}"
            
            # Формат ДДММ.ГГ (пропущена точка между днем и месяцем)
            match = re.match(r'^(\d{2})(\d{2})\.(\d{2})$', date_str)
            if match:
                day, month, year = match.groups()
                full_year = self._expand_year(year)
                return f"{int(day):02d}.{int(month):02d}.{full_year}"
            
            # Формат ДДММГГГГ (без разделителей)
            match = re.match(r'^(\d{2})(\d{2})(\d{4})$', date_str)
            if match:
                day, month, year = match.groups()
                return f"{int(day):02d}.{int(month):02d}.{year}"
            
            # Формат ДД/ММ/ГГ
            match = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{2})$', date_str)
            if match:
                day, month, year = match.groups()
                full_year = self._expand_year(year)
                return f"{int(day):02d}.{int(month):02d}.{full_year}"
            
            # Формат ДД/ММ/ГГГГ
            match = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{4})$', date_str)
            if match:
                day, month, year = match.groups()
                return f"{int(day):02d}.{int(month):02d}.{year}"
            
            # Формат ДД-ММ-ГГ
            match = re.match(r'^(\d{1,2})-(\d{1,2})-(\d{2})$', date_str)
            if match:
                day, month, year = match.groups()
                full_year = self._expand_year(year)
                return f"{int(day):02d}.{int(month):02d}.{full_year}"
            
            # Формат ДД-ММ-ГГГГ
            match = re.match(r'^(\d{1,2})-(\d{1,2})-(\d{4})$', date_str)
            if match:
                day, month, year = match.groups()
                return f"{int(day):02d}.{int(month):02d}.{year}"
            
            # Если ни один из форматов не подошел, возвращаем None
            return None
    
    def _expand_year(self, year_str):
        """
        Преобразует двузначный год в четырехзначный
        Правило: 00-25 -> 2000-2025, 26-99 -> 1926-1999
        """
        year = int(year_str)
        
        # Определяем текущий год для расчета порога преобразования
        current_year = datetime.now().year
        current_short_year = current_year % 100
        
        if year <= current_short_year:
            return 2000 + year
        else:
            return 1900 + year
    
    def _is_date(self, value):
        """
        Проверяет, является ли значение ячейки датой
        """
        if not value:
            return False
            
        # Преобразуем значение в строку и удаляем пробелы
        value_str = str(value).strip()
        
        # Шаблоны для проверки дат в различных форматах
        date_patterns = [
            r'^\d{1,2}\.\d{1,2}\.\d{2,4}',  # ДД.ММ.ГГ или ДД.ММ.ГГГГ
            r'^\d{1,2}/\d{1,2}/\d{2,4}',    # ДД/ММ/ГГ или ДД/ММ/ГГГГ
            r'^\d{1,2}-\d{1,2}-\d{2,4}',    # ДД-ММ-ГГ или ДД-ММ-ГГГГ
            r'^\d{2}\d{2}\.\d{2}'           # ДДММ.ГГ (пропущена точка между днем и месяцем)
        ]
        
        for pattern in date_patterns:
            if re.match(pattern, value_str):
                return True
                
        return False

    def _adjust_column_width(self, sheet):
        """Автоподбор ширины столбцов"""
        for column_cells in sheet.columns:
            max_length = 0
            column = column_cells[0].column_letter
            for cell in column_cells:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            sheet.column_dimensions[column].width = adjusted_width
    
    def open_file(self, file_path):
        """Открытие файла в соответствующем приложении"""
        try:
            if platform.system() == 'Windows':
                os.startfile(file_path)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', file_path])
            else:  # Linux и другие
                subprocess.call(['xdg-open', file_path])
            return True
        except Exception as e:
            print(f"Ошибка при открытии файла: {str(e)}")
            return False


if __name__ == "__main__":
    root = tk.Tk()
    app = SimpleDocxToExcelApp(root)
    root.mainloop()